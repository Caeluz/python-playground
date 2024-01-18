import os
import fitz  # PyMuPDF
from docx import Document

def extract_text_with_style(pdf_path):
    pdf_document = fitz.open(pdf_path)
    doc = Document()

    for page_number in range(pdf_document.page_count):
        page = pdf_document[page_number]
        
        # Iterate through the page's text blocks
        for block in page.get_text("blocks"):
            text = block[4]  + " " # Extract the text
            # font_size = block[1]  # Font size
            # font_bold = block[6] == 1 
            
            # Add a new run for each block of text
        # run = doc.paragraphs[-1].add_run(text)

            # if font_size > 14.4 and font_bold:
            #     # Large Bold Text
            #     doc.add_paragraph(text, style='Heading1')
            # else:
            #     # Paragraph-form Text
        doc.add_paragraph(text)

    pdf_document.close()

    # Save the DOCX file in the same directory as the input PDF
    pdf_directory = os.path.dirname(pdf_path)
    docx_filename = os.path.splitext(os.path.basename(pdf_path))[0] + "_output.docx"
    docx_path = os.path.join(pdf_directory, docx_filename)
    doc.save(docx_path)

if __name__ == "__main__":
    pdf_path = r'C:\Users\ADMIN\Documents\lgu\System\LGU MAGAZINE.pdf'
    extract_text_with_style(pdf_path)
